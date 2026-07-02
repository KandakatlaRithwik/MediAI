"""
Report extraction engine (Module 3, Step 1).

Extracts raw text from an uploaded medical report (PDF or TXT), preserving
line structure so ReportParserService's regex patterns can reliably find
"Label: Value" style lines.

PDF strategy: PyPDFLoader first (consistent with the rest of the codebase's
document-loading approach, app.rag.document_loader). Some lab-report PDFs
are malformed or use layouts PyPDFLoader's underlying parser struggles with
(e.g. table-heavy reports), so on failure - or if PyPDFLoader returns no
extractable text - this falls back to pdfplumber, which uses a different
underlying parser and handles many of those cases PyPDFLoader cannot.

Image support (scanned reports / OCR) is explicitly out of scope for now
(see Future Ready note in the spec) - extract_text() raises
DocumentProcessingError for image file types so the route can give a clear
"not yet supported" message rather than silently failing.
"""

import logging
from pathlib import Path

import pdfplumber
from langchain_community.document_loaders import PyPDFLoader, TextLoader

from app.core.exceptions import DocumentProcessingError

logger = logging.getLogger("report_analysis")

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}


def _extract_pdf_with_pypdfloader(file_path: str) -> str:
    loader = PyPDFLoader(file_path)
    documents = loader.load()
    return "\n".join(doc.page_content for doc in documents if doc.page_content)


def _extract_pdf_with_pdfplumber(file_path: str) -> str:
    pages_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
    return "\n".join(pages_text)


def _extract_pdf(file_path: str) -> str:
    """Extract text from a PDF, trying PyPDFLoader first and falling back
    to pdfplumber if it fails or yields no usable text."""
    try:
        text = _extract_pdf_with_pypdfloader(file_path)
        if text.strip():
            logger.info("Extracted PDF text via PyPDFLoader (%d chars): %s", len(text), file_path)
            return text
        logger.warning("PyPDFLoader returned no text for '%s'; falling back to pdfplumber.", file_path)
    except Exception as exc:
        logger.warning("PyPDFLoader failed for '%s' (%s); falling back to pdfplumber.", file_path, exc)

    try:
        text = _extract_pdf_with_pdfplumber(file_path)
    except Exception as exc:
        logger.error("pdfplumber fallback also failed for '%s': %s", file_path, exc)
        raise DocumentProcessingError(f"Failed to extract text from PDF report: {exc}") from exc

    if not text.strip():
        raise DocumentProcessingError(
            "No extractable text found in this PDF report. It may be a scanned/image-only "
            "document, which is not yet supported."
        )

    logger.info("Extracted PDF text via pdfplumber fallback (%d chars): %s", len(text), file_path)
    return text


def _extract_txt(file_path: str) -> str:
    try:
        loader = TextLoader(file_path, encoding="utf-8")
        documents = loader.load()
    except Exception as exc:
        logger.error("Failed to read text report '%s': %s", file_path, exc)
        raise DocumentProcessingError(f"Failed to read text report: {exc}") from exc

    text = documents[0].page_content if documents else ""
    if not text.strip():
        raise DocumentProcessingError("The uploaded text report is empty or unreadable.")

    logger.info("Extracted TXT report text (%d chars): %s", len(text), file_path)
    return text


def _extract_docx(file_path: str) -> str:
    """Extract raw text from a DOCX using python-docx.

    Both paragraph text and table cells are captured so lab reports laid
    out as tables (very common) still surface each parameter/value pair
    on its own line for the downstream parser.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise DocumentProcessingError(
            "DOCX support requires the 'python-docx' package to be installed."
        ) from exc

    try:
        document = Document(file_path)
    except Exception as exc:
        logger.error("Failed to open DOCX '%s': %s", file_path, exc)
        raise DocumentProcessingError(f"Failed to read DOCX report: {exc}") from exc

    lines = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                lines.append(" \t ".join(cells))

    text = "\n".join(lines)
    if not text.strip():
        raise DocumentProcessingError("The uploaded DOCX report is empty or unreadable.")

    logger.info("Extracted DOCX text (%d chars): %s", len(text), file_path)
    return text


class ReportExtractionService:
    def extract_text(self, file_path: str) -> str:
        """Extract raw text from a medical report file (PDF, TXT, DOCX).

        Image-based reports are handled separately by ReportImageParserService
        via the /analyze-image-report endpoint.
        """
        extension = Path(file_path).suffix.lower()
        logger.info("Extraction: file=%s ext=%s", file_path, extension)

        if extension in IMAGE_EXTENSIONS:
            raise DocumentProcessingError(
                "Image-based reports must be uploaded via the /analyze-image-report endpoint "
                "(the frontend routes JPG/PNG uploads there automatically)."
            )
        if extension == ".pdf":
            return _extract_pdf(file_path)
        if extension == ".txt":
            return _extract_txt(file_path)
        if extension in (".docx",):
            return _extract_docx(file_path)

        raise DocumentProcessingError(f"Unsupported report file extension: '{extension}'.")
